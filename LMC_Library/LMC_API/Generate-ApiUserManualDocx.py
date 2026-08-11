#!/usr/bin/env python3
"""Generate an editable Word source for the LASAL Motion Control API manual.

The Markdown file is the shared content source for the PDF and DOCX manuals.
Tables, headings, lists, and code listings are emitted as native Word objects so
that a recipient can edit the document without the internal build environment.
"""

from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


BLUE = "123B66"
MID_BLUE = "26669A"
ACCENT = "00A6D6"
LIGHT_BLUE = "EAF5FB"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "66717C"
DARK = "20262C"
BORDER = "CBD4DC"
CODE_BG = "F6F8FA"
WARNING_BG = "FFF7E6"
WARNING_BORDER = "D58A00"
WHITE = "FFFFFF"

BODY_FONT = "Malgun Gothic"
CODE_FONT = "Consolas"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(child)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int = 90, start: int = 110,
                     bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start),
                          ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: int = 5) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    tbl_pr.append(borders)


def set_east_asia(run, name: str) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def set_style_font(style, name: str, size: float, color: str | None = None,
                   bold: bool | None = None) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.append(begin)
    run.append(instr)
    run.append(separate)
    run.append(text)
    run.append(end)


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    current = settings.find(qn("w:updateFields"))
    if current is None:
        current = OxmlElement("w:updateFields")
        settings.append(current)
    current.set(qn("w:val"), "true")


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    set_style_font(normal, BODY_FONT, 9.3, DARK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.22

    heading_1 = styles["Heading 1"]
    set_style_font(heading_1, BODY_FONT, 18, BLUE, True)
    heading_1.paragraph_format.space_before = Pt(10)
    heading_1.paragraph_format.space_after = Pt(6)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    set_style_font(heading_2, BODY_FONT, 12.5, MID_BLUE, True)
    heading_2.paragraph_format.space_before = Pt(9)
    heading_2.paragraph_format.space_after = Pt(4)
    heading_2.paragraph_format.keep_with_next = True

    heading_3 = styles["Heading 3"]
    set_style_font(heading_3, BODY_FONT, 10.5, MID_BLUE, True)
    heading_3.paragraph_format.space_before = Pt(7)
    heading_3.paragraph_format.space_after = Pt(3)
    heading_3.paragraph_format.keep_with_next = True

    toc_heading = styles.add_style("Manual TOC Heading", 1)
    set_style_font(toc_heading, BODY_FONT, 19, BLUE, True)
    toc_heading.paragraph_format.space_after = Pt(12)
    toc_heading.paragraph_format.page_break_before = True

    # Word expands the TOC field with the built-in TOC 1/TOC 2 paragraph
    # styles.  Pin compact styles here so a normal field refresh does not leave
    # a single orphan entry on a third TOC page.
    for name, size, indent, bold in (
        ("TOC 1", 9.5, 0, True),
        ("TOC 2", 8.8, 4, False),
    ):
        try:
            toc_style = styles[name]
        except KeyError:
            toc_style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style_element = toc_style._element
        style_element.attrib.pop(qn("w:customStyle"), None)
        style_name = style_element.find(qn("w:name"))
        if style_name is not None:
            style_name.set(qn("w:val"), name.lower())
        # CT_Style child order is schema-significant.  Insert each new
        # metadata element before its first legal successor instead of using a
        # fixed index, which would reverse the nodes and leave an invalid DOCX
        # until Word happened to repair it on save.
        metadata = (
            ("basedOn", "Normal", ("next", "link", "autoRedefine", "hidden", "uiPriority", "semiHidden", "unhideWhenUsed", "qFormat", "locked", "personal", "personalCompose", "personalReply", "rsid", "pPr", "rPr")),
            ("next", "Normal", ("link", "autoRedefine", "hidden", "uiPriority", "semiHidden", "unhideWhenUsed", "qFormat", "locked", "personal", "personalCompose", "personalReply", "rsid", "pPr", "rPr")),
            ("uiPriority", "39", ("semiHidden", "unhideWhenUsed", "qFormat", "locked", "personal", "personalCompose", "personalReply", "rsid", "pPr", "rPr")),
            ("semiHidden", None, ("unhideWhenUsed", "qFormat", "locked", "personal", "personalCompose", "personalReply", "rsid", "pPr", "rPr")),
            ("unhideWhenUsed", None, ("qFormat", "locked", "personal", "personalCompose", "personalReply", "rsid", "pPr", "rPr")),
        )
        for tag, value, successors in metadata:
            node = style_element.find(qn(f"w:{tag}"))
            if node is None:
                node = OxmlElement(f"w:{tag}")
                style_element.insert_element_before(
                    node,
                    *(f"w:{successor}" for successor in successors),
                )
            if value is not None:
                node.set(qn("w:val"), value)
        set_style_font(toc_style, BODY_FONT, size, DARK, bold)
        toc_style.paragraph_format.left_indent = Mm(indent)
        toc_style.paragraph_format.space_before = Pt(0)
        toc_style.paragraph_format.space_after = Pt(0)
        toc_style.paragraph_format.line_spacing = 1.0

    note = styles.add_style("Manual Note", 1)
    set_style_font(note, BODY_FONT, 8, MID_GRAY)
    note.paragraph_format.space_after = Pt(5)

    table_cell = styles.add_style("Manual Table Cell", 1)
    set_style_font(table_cell, BODY_FONT, 7.7, DARK)
    table_cell.paragraph_format.space_after = Pt(0)
    table_cell.paragraph_format.line_spacing = 1.05

    code = styles.add_style("Manual Code", 1)
    set_style_font(code, CODE_FONT, 7.2, "17202A")
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0

    callout = styles.add_style("Manual Callout", 1)
    set_style_font(callout, BODY_FONT, 8.5, DARK)
    callout.paragraph_format.space_after = Pt(0)
    callout.paragraph_format.line_spacing = 1.15


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(18)
    section.right_margin = Mm(16)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(17)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)
    section.different_first_page_header_footer = True

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    left = p.add_run("LASAL MOTION CONTROL API")
    set_east_asia(left, BODY_FONT)
    left.font.size = Pt(7.3)
    left.font.bold = True
    left.font.color.rgb = RGBColor.from_string(BLUE)
    right = p.add_run("    API GUIDE")
    set_east_asia(right, BODY_FONT)
    right.font.size = Pt(7.3)
    right.font.color.rgb = RGBColor.from_string(MID_GRAY)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    version = p.add_run("LasalMotionControlLib 0.9.1-preview    ")
    set_east_asia(version, BODY_FONT)
    version.font.size = Pt(7)
    version.font.color.rgb = RGBColor.from_string(MID_GRAY)
    add_field(p, " PAGE ", "1")

    set_update_fields(document)


def parse_cover(lines: list[str]) -> tuple[str, dict[str, str]]:
    title = "LASAL Motion Control API 기능 설명서"
    metadata: dict[str, str] = {}
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            title = stripped[2:].strip()
        match = re.match(r"^([^:]+):\s*(.+?)(?:\s{2})?$", line.rstrip())
        if match:
            metadata[match.group(1).strip()] = match.group(2).strip()
    return title, metadata


def add_cover(document: Document, title: str, metadata: dict[str, str]) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(25)

    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = False
    banner.columns[0].width = Cm(17.6)
    cell = banner.cell(0, 0)
    cell.width = Cm(17.6)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, BLUE)
    set_cell_margins(cell, top=420, start=520, bottom=420, end=520)
    set_table_borders(banner, BLUE, 0)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("LASAL")
    set_east_asia(run, BODY_FONT)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title.replace(" API ", " API\n"))
    set_east_asia(run, BODY_FONT)
    run.font.size = Pt(27)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(WHITE)

    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        "LasalMotionControlLib 공개 C# API의\n"
        "기능, 호출 인자 UNIT과 반환값"
    )
    set_east_asia(run, BODY_FONT)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string("DCECF5")

    document.add_paragraph().paragraph_format.space_after = Pt(20)

    meta = document.add_table(rows=0, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    meta.columns[0].width = Cm(3.4)
    meta.columns[1].width = Cm(11.8)
    labels = (
        ("문서", "API 기능 설명서"),
        ("적용 API", metadata.get("적용 API", "LasalMotionControlLib 0.9.1-preview")),
        ("환경", metadata.get("대상 환경", "Windows, .NET Framework 4.8")),
        ("발행", metadata.get("발행일", "2026-07-16")),
        ("문서 버전", metadata.get("문서 버전", "1.2")),
    )
    for label, value in labels:
        cells = meta.add_row().cells
        for item in cells:
            set_cell_margins(item, top=100, start=130, bottom=100, end=130)
        set_cell_shading(cells[0], LIGHT_BLUE)
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_east_asia(run, BODY_FONT)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(BLUE)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_east_asia(run, BODY_FONT)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(DARK)
    set_table_borders(meta)

    note = document.add_paragraph(style="Manual Note")
    note.paragraph_format.space_before = Pt(13)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "API 기능과 호출에 필요한 값만 간단히 설명합니다."
    )
    note.add_run().add_break(WD_BREAK.PAGE)


INLINE_PATTERN = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")


def add_inline(paragraph, text: str, *, base_size: float | None = None,
               base_color: str | None = None, bold: bool = False) -> None:
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        is_code = part.startswith("`") and part.endswith("`")
        is_bold = part.startswith("**") and part.endswith("**")
        value = part[1:-1] if is_code else part[2:-2] if is_bold else part
        run = paragraph.add_run(value)
        set_east_asia(run, CODE_FONT if is_code else BODY_FONT)
        if base_size is not None:
            run.font.size = Pt(base_size - 0.2 if is_code else base_size)
        if base_color:
            run.font.color.rgb = RGBColor.from_string(base_color)
        if is_code:
            run.font.color.rgb = RGBColor.from_string("1C4E73")
        run.font.bold = bold or is_bold


def column_factors(count: int) -> list[float]:
    if count == 1:
        return [1.0]
    if count == 2:
        return [0.42, 0.58]
    if count == 3:
        return [0.25, 0.21, 0.54]
    if count == 4:
        return [0.19, 0.18, 0.23, 0.40]
    return [1.0 / count] * count


def add_native_table(document: Document, rows: list[list[str]]) -> None:
    count = max(len(row) for row in rows)
    normalized = [row + [""] * (count - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    factors = column_factors(count)
    content_width = 17.6
    for index, factor in enumerate(factors):
        table.columns[index].width = Cm(content_width * factor)
    for row_index, row in enumerate(normalized):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.style = document.styles["Manual Table Cell"]
            if row_index == 0 or (
                len(normalized) <= 6 and row_index < len(normalized) - 1
            ):
                # Keep the header with at least one data row. Small reference
                # tables are kept together when the remaining page permits it.
                p.paragraph_format.keep_with_next = True
            add_inline(
                p,
                value,
                base_size=7.7,
                base_color=WHITE if row_index == 0 else DARK,
                bold=row_index == 0,
            )
    set_repeat_table_header(table.rows[0])
    set_table_borders(table)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def add_code_block(document: Document, code: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(17.6)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_BG)
    set_cell_margins(cell, top=110, start=140, bottom=110, end=140)
    set_table_borders(table, BORDER, 5)
    p = cell.paragraphs[0]
    p.style = document.styles["Manual Code"]
    lines = code.rstrip("\n").splitlines() or [""]
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_east_asia(run, CODE_FONT)
        run.font.size = Pt(7.2)
        run.font.color.rgb = RGBColor.from_string("17202A")
        if index != len(lines) - 1:
            run.add_break()
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def add_callout(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(17.6)
    cell = table.cell(0, 0)
    set_cell_shading(cell, WARNING_BG)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    set_table_borders(table, WARNING_BORDER, 7)
    p = cell.paragraphs[0]
    p.style = document.styles["Manual Callout"]
    add_inline(p, text, base_size=8.5, base_color=DARK)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        row = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in row):
            rows.append(row)
        index += 1
    return rows, index


def add_toc(document: Document) -> None:
    document.add_paragraph("목차", style="Manual TOC Heading")
    p = document.add_paragraph()
    add_field(p, ' TOC \\o "1-2" \\h \\z \\u ', "목차 업데이트 필요")
    document.add_page_break()


def add_body(document: Document, source: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    try:
        first_break = lines.index(r"\pagebreak")
    except ValueError as exc:
        raise ValueError("Manual source must contain an initial \\pagebreak") from exc
    title, metadata = parse_cover(lines[:first_break])
    add_cover(document, title, metadata)

    index = first_break + 1
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(part.strip() for part in paragraph_buffer).strip()
        paragraph_buffer.clear()
        if text:
            p = document.add_paragraph()
            add_inline(p, text)

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index >= len(lines):
                raise ValueError("Unclosed code fence in manual source")
            add_code_block(document, "\n".join(code_lines))
            index += 1
            continue

        if stripped == r"\pagebreak":
            flush_paragraph()
            document.add_page_break()
            index += 1
            continue

        if stripped == r"\toc":
            flush_paragraph()
            add_toc(document)
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            callout_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                callout_lines.append(
                    re.sub(r"^>\s?", "", lines[index].strip()).strip()
                )
                index += 1
            add_callout(document, " ".join(callout_lines))
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            p = document.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading.group(2).strip())
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            rows, index = parse_table(lines, index)
            if rows:
                add_native_table(document, rows)
            continue

        if re.match(r"^\s*-\s+", raw):
            flush_paragraph()
            while index < len(lines) and re.match(r"^\s*-\s+", lines[index]):
                value = re.sub(r"^\s*-\s+", "", lines[index]).strip()
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Mm(5)
                p.paragraph_format.first_line_indent = Mm(-4)
                marker = p.add_run("•\t")
                set_east_asia(marker, BODY_FONT)
                marker.font.size = Pt(9.3)
                marker.font.color.rgb = RGBColor.from_string(DARK)
                add_inline(p, value)
                index += 1
            continue

        if re.match(r"^\s*\d+\.\s+", raw):
            flush_paragraph()
            while index < len(lines) and re.match(r"^\s*\d+\.\s+", lines[index]):
                match = re.match(r"^\s*(\d+)\.\s+(.+)$", lines[index])
                if match is None:
                    break
                number = match.group(1)
                value = match.group(2).strip()
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Mm(7)
                p.paragraph_format.first_line_indent = Mm(-7)
                marker = p.add_run(number + ".\t")
                set_east_asia(marker, BODY_FONT)
                marker.font.size = Pt(9.3)
                marker.font.color.rgb = RGBColor.from_string(DARK)
                add_inline(p, value)
                index += 1
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source",
        type=Path,
        default=Path(__file__).with_name("API_USER_MANUAL_KO.md"),
    )
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--distribution-copy", type=Path)
    args = parser.parse_args()

    source = args.source.resolve()
    if not source.exists():
        raise FileNotFoundError(f"Manual source not found: {source}")

    document = Document()
    configure_styles(document)
    configure_document(document)
    add_body(document, source)

    properties = document.core_properties
    properties.title = "LASAL Motion Control API 기능 설명서"
    properties.author = "Elmo Motion Control / LASAL API Project"
    properties.subject = "LasalMotionControlLib API usage"
    properties.comments = "Editable source generated from API_USER_MANUAL_KO.md"

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)

    if args.distribution_copy:
        args.distribution_copy.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(args.output, args.distribution_copy)

    print(f"Generated: {args.output.resolve()}")
    if args.distribution_copy:
        print(f"Copied: {args.distribution_copy.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
